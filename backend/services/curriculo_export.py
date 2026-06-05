from datetime import datetime
from typing import Any

def _clean_text(text: str) -> str:
    if not text:
        return ""
    replacements = {
        "\u25cf": "-", "\u2022": "-", "\u2013": "-", "\u2014": "--",
        "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"', "\u2026": "..."
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.encode("latin-1", errors="replace").decode("latin-1")

def generate_pdf(doc: dict, versao_label: str = "") -> bytes:
    from fpdf import FPDF
    
    estruturado = doc.get("estruturado") or doc
    
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font("Helvetica", "", 10)

    def section_title(text: str):
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_draw_color(180, 180, 180)
        pdf.line(pdf.get_x(), pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)
        pdf.cell(0, 6, _clean_text(text).upper(), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)

    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 10, _clean_text(estruturado.get("nome", "")), align="C", new_x="LMARGIN", new_y="NEXT")
    
    contato_parts = []
    if estruturado.get("email"): contato_parts.append(estruturado["email"])
    if estruturado.get("telefone"): contato_parts.append(estruturado["telefone"])
    if estruturado.get("cidade"): contato_parts.append(estruturado["cidade"])
    
    if contato_parts:
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 5, _clean_text(" | ".join(contato_parts)), align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        
    socials = []
    if estruturado.get("linkedin"): socials.append(f"LinkedIn: {estruturado['linkedin']}")
    if estruturado.get("github"): socials.append(f"GitHub: {estruturado['github']}")
    if estruturado.get("portfolio"): socials.append(f"Portfolio: {estruturado['portfolio']}")
    
    if socials:
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(0, 5, _clean_text(" | ".join(socials)), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    if estruturado.get("resumo_profissional"):
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5.5, _clean_text(estruturado["resumo_profissional"]), align="J")
        pdf.ln(2)

    if estruturado.get("experiencias"):
        section_title("EXPERIENCIA PROFISSIONAL")
        for exp in estruturado["experiencias"]:
            pdf.set_font("Helvetica", "B", 10)
            cargo = exp.get("cargo", "")
            empresa = exp.get("empresa", "")
            periodo = f" ({exp.get('data_inicio', '')} - {exp.get('data_fim', '')})" if exp.get('data_inicio') else ""
            pdf.cell(0, 6, _clean_text(f"{cargo} | {empresa}{periodo}"), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 9)
            for b in exp.get("descricao", []):
                if b.strip():
                    pdf.set_x(pdf.l_margin + 5)
                    pdf.multi_cell(0, 5, _clean_text(f"- {b}"))
            pdf.ln(2)

    if estruturado.get("projetos"):
        section_title("PROJETOS")
        for proj in estruturado["projetos"]:
            pdf.set_font("Helvetica", "B", 10)
            nome = proj.get("nome", "")
            stack = f" [{', '.join(proj.get('tecnologias', []))}]" if proj.get("tecnologias") else ""
            pdf.cell(0, 6, _clean_text(f"{nome}{stack}"), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 9)
            if proj.get("descricao"):
                pdf.multi_cell(0, 5, _clean_text(proj["descricao"]))
            for b in proj.get("bullets", []):
                if b.strip():
                    pdf.set_x(pdf.l_margin + 5)
                    pdf.multi_cell(0, 5, _clean_text(f"- {b}"))
            pdf.ln(1)

    if estruturado.get("skills"):
        section_title("COMPETENCIAS TECNICAS")
        pdf.set_font("Helvetica", "", 9)
        skills_list = []
        for s in estruturado["skills"]:
            if isinstance(s, dict):
                skills_list.append(s.get("nome", ""))
            else:
                skills_list.append(str(s))
        pdf.multi_cell(0, 5, _clean_text("  *  ".join(skills_list)))
        pdf.ln(2)

    if estruturado.get("formacoes"):
        section_title("FORMACAO ACADEMICA")
        for f in estruturado["formacoes"]:
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 6, _clean_text(f.get("curso", "")), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 9)
            inst = f.get("instituicao", "")
            data = f.get("data_conclusao", "")
            nivel = f.get("nivel", "")
            detalhes = " - ".join(x for x in [nivel, data] if x)
            if inst or detalhes:
                pdf.cell(0, 5, _clean_text(f"{inst}{' - ' + detalhes if detalhes else ''}"), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)

    if estruturado.get("certificacoes"):
        section_title("CERTIFICACOES")
        pdf.set_font("Helvetica", "", 10)
        for c in estruturado["certificacoes"]:
            partes = ", ".join(x for x in [c.get("nome", ""), c.get("instituicao", ""), c.get("ano", "")] if x)
            pdf.multi_cell(0, 6, _clean_text(partes))
        pdf.ln(1)

    if estruturado.get("idiomas"):
        section_title("IDIOMAS")
        pdf.set_font("Helvetica", "", 10)
        for i in estruturado["idiomas"]:
            pdf.cell(0, 6, _clean_text(f"{i.get('idioma', '')}: {i.get('nivel', '')}"), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)

    pdf.ln(5)
    pdf.set_font("Helvetica", "", 7)
    pdf.set_text_color(150, 150, 150)
    footer = f"Curriculo gerado pela Work Job Platform | Versao: {versao_label or 'atual'} | {datetime.utcnow().strftime('%d/%m/%Y %H:%M')} UTC"
    pdf.cell(0, 4, _clean_text(footer), align="C", new_x="LMARGIN", new_y="NEXT")

    return bytes(pdf.output())

def generate_docx(doc: dict, versao_label: str = "") -> bytes:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    estruturado = doc.get("estruturado") or doc

    docx = Document()
    style = docx.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    p = docx.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(estruturado.get("nome", ""))
    run.bold = True
    run.font.size = Pt(22)
    
    contato_parts = []
    if estruturado.get("email"): contato_parts.append(estruturado["email"])
    if estruturado.get("telefone"): contato_parts.append(estruturado["telefone"])
    if estruturado.get("cidade"): contato_parts.append(estruturado["cidade"])
    
    if contato_parts:
        p2 = docx.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(" | ".join(contato_parts))
        run2.font.size = Pt(9)
        
    socials = []
    if estruturado.get("linkedin"): socials.append(f"LinkedIn: {estruturado['linkedin']}")
    if estruturado.get("github"): socials.append(f"GitHub: {estruturado['github']}")
    if estruturado.get("portfolio"): socials.append(f"Portfolio: {estruturado['portfolio']}")
    
    if socials:
        p3 = docx.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(" | ".join(socials))
        run3.font.size = Pt(9)

    if estruturado.get("resumo_profissional"):
        docx.add_paragraph(estruturado["resumo_profissional"])

    if estruturado.get("experiencias"):
        docx.add_heading("Experiencia Profissional", level=2)
        for exp in estruturado["experiencias"]:
            cargo = exp.get("cargo", "")
            empresa = exp.get("empresa", "")
            periodo = f" ({exp.get('data_inicio', '')} - {exp.get('data_fim', '')})" if exp.get('data_inicio') else ""
            p = docx.add_paragraph()
            p.add_run(f"{cargo} | {empresa}{periodo}").bold = True
            for b in exp.get("descricao", []):
                if b.strip():
                    docx.add_paragraph(b, style="List Bullet")

    if estruturado.get("projetos"):
        docx.add_heading("Projetos", level=2)
        for proj in estruturado["projetos"]:
            nome = proj.get("nome", "")
            stack = f" [{', '.join(proj.get('tecnologias', []))}]" if proj.get("tecnologias") else ""
            p = docx.add_paragraph()
            p.add_run(f"{nome}{stack}").bold = True
            if proj.get("descricao"):
                docx.add_paragraph(proj["descricao"])
            for b in proj.get("bullets", []):
                if b.strip():
                    docx.add_paragraph(b, style="List Bullet")

    if estruturado.get("skills"):
        docx.add_heading("Competencias Tecnicas", level=2)
        skills_list = []
        for s in estruturado["skills"]:
            if isinstance(s, dict):
                skills_list.append(s.get("nome", ""))
            else:
                skills_list.append(str(s))
        docx.add_paragraph(" * ".join(skills_list))

    if estruturado.get("formacoes"):
        docx.add_heading("Formacao Academica", level=2)
        for f in estruturado["formacoes"]:
            p = docx.add_paragraph()
            p.add_run(f.get("curso", "")).bold = True
            inst = f.get("instituicao", "")
            data = f.get("data_conclusao", "")
            nivel = f.get("nivel", "")
            detalhes = ", ".join(x for x in [nivel, data] if x)
            if inst or detalhes:
                docx.add_paragraph(f"{inst} - {detalhes}".lstrip(" - "))

    if estruturado.get("certificacoes"):
        docx.add_heading("Certificacoes", level=2)
        for c in estruturado["certificacoes"]:
            partes = ", ".join(x for x in [c.get("nome", ""), c.get("instituicao", ""), c.get("ano", "")] if x)
            docx.add_paragraph(partes)

    if estruturado.get("idiomas"):
        docx.add_heading("Idiomas", level=2)
        for i in estruturado["idiomas"]:
            docx.add_paragraph(f"{i.get('idioma', '')}: {i.get('nivel', '')}")

    footer_text = f"Curriculo gerado pela Work Job Platform | Versao: {versao_label or 'atual'}"
    p = docx.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(footer_text)
    run.font.size = Pt(8)

    buf = io.BytesIO()
    docx.save(buf)
    buf.seek(0)
    return buf.read()
